import pytest
from docx import Document

from ppt2course.script_extract import ScriptExtractionError, extract_text_from_file


def test_extracts_text_from_plain_txt_bytes():
    content = "第1頁\n大家好\n第2頁\n謝謝收看".encode("utf-8")
    result = extract_text_from_file("script.txt", content)
    assert result == "第1頁\n大家好\n第2頁\n謝謝收看"


def test_strips_leading_utf8_bom_from_txt():
    # a BOM-prefixed .txt (common from Windows editors saving "UTF-8 with
    # BOM") left a stray ﻿ before the first line's text, which broke
    # the "^第" page-marker regex on line 1 only — reported symptom: page 1
    # wasn't recognized unless a blank line was added before "第1頁".
    content = "﻿第1頁\n大家好".encode("utf-8")
    result = extract_text_from_file("script.txt", content)
    assert result == "第1頁\n大家好"
    assert not result.startswith("﻿")


def test_extracts_paragraphs_from_docx_joined_by_newline(tmp_path):
    doc = Document()
    doc.add_paragraph("第1頁")
    doc.add_paragraph("大家好")
    doc.add_paragraph("第2頁")
    doc.add_paragraph("謝謝收看")
    docx_path = tmp_path / "script.docx"
    doc.save(docx_path)

    result = extract_text_from_file("script.docx", docx_path.read_bytes())
    assert result == "第1頁\n大家好\n第2頁\n謝謝收看"


def test_docx_skips_fully_blank_paragraphs(tmp_path):
    doc = Document()
    doc.add_paragraph("第1頁")
    doc.add_paragraph("")
    doc.add_paragraph("大家好")
    docx_path = tmp_path / "script.docx"
    doc.save(docx_path)

    result = extract_text_from_file("script.docx", docx_path.read_bytes())
    assert result == "第1頁\n大家好"


def test_unsupported_extension_raises():
    with pytest.raises(ScriptExtractionError):
        extract_text_from_file("script.pdf", b"whatever")


def test_corrupt_docx_raises():
    with pytest.raises(ScriptExtractionError):
        extract_text_from_file("script.docx", b"not a real docx file")


def test_extension_matching_is_case_insensitive():
    content = "第1頁\n內容".encode("utf-8")
    result = extract_text_from_file("SCRIPT.TXT", content)
    assert result == "第1頁\n內容"
