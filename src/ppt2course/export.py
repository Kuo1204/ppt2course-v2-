"""STEP 6: Export — write 課程.mp4 / 字幕.srt / 講稿.docx with a shared base filename."""

import os
import shutil

from docx import Document


class ExportError(Exception):
    pass


def _build_script_document(scripts: list[str]) -> Document:
    doc = Document()
    for i, script in enumerate(scripts, start=1):
        doc.add_heading(f"投影片 {i}", level=1)
        doc.add_paragraph(script)
    return doc


def export_outputs(
    mp4_source_path: str,
    srt_source_path: str,
    scripts: list[str],
    out_dir: str,
    base_name: str,
) -> dict[str, str]:
    os.makedirs(out_dir, exist_ok=True)

    mp4_dest = os.path.join(out_dir, f"{base_name}.mp4")
    srt_dest = os.path.join(out_dir, f"{base_name}.srt")
    docx_dest = os.path.join(out_dir, f"{base_name}.docx")

    for dest in (mp4_dest, srt_dest, docx_dest):
        if os.path.exists(dest):
            raise ExportError(f"output file already exists: {dest}")

    try:
        shutil.copyfile(mp4_source_path, mp4_dest)
        shutil.copyfile(srt_source_path, srt_dest)
        _build_script_document(scripts).save(docx_dest)
    except OSError as exc:
        raise ExportError(f"failed to export outputs: {exc}") from exc

    return {"mp4": mp4_dest, "srt": srt_dest, "docx": docx_dest}
